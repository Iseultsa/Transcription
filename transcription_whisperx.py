"""
Transcription_IA.py — Script de transcription audio avec WhisperX et diarisation automatique

Autrice : Iseult Séguin Aubé
Date : 2025-07-25

Licence : BSD 3-Clause (voir fichier LICENSE)

Ce script a été développé dans un contexte de recherche et permet de transcrire un fichier audio
en texte, avec identification automatique des locuteurs (diarisation). Il utilise :

- WhisperX (Apache 2.0), une extension de Whisper pour l'alignement et la gestion avancée des segments.
- pyannote.audio (MIT), pour la diarisation automatique des locuteurs via des modèles préentraînés.
- python-docx (MIT), pour l’exportation facultative en format Word (.docx).
- Les bibliothèques standard de Python (os, datetime, etc.).

📌Exigences système et environnement :

- Python 3.8 ou plus récent
- GPU compatible CUDA (obligatoire, aucun fallback CPU)
- NVIDIA CUDA Toolkit et drivers installés
- whisperx :     pip install git+https://github.com/m-bain/whisperx.git
- pyannote.audio : pip install pyannote.audio
- Clé d’API Hugging Face (obligatoire pour la diarisation) enregistrée dans le cache local :
    → https://huggingface.co/settings/tokens
- python-docx (facultatif pour l’export Word) : pip install python-docx

❗ Ce script ne fonctionne que sur une machine équipée de GPU + CUDA.
❗ La diarisation est toujours activée, donc pyannote.audio et une clé Hugging Face sont requis.

Aucune garantie de maintenance n’est assurée. Ce script est fourni "en l'état", à des fins de transparence
scientifique et de réutilisation éventuelle.

Conditions d'utilisation :
- La licence BSD 3-Clause permet la réutilisation libre du code, à condition de ne pas utiliser le nom
  de l’autrice (Iseult Séguin Aubé) pour faire la promotion de produits dérivés sans autorisation explicite.
- Les utisatrices et ulitilisateurs sont responsables de respecter les licences des dépendances citées.

"""

import whisperx
import os
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️  Module python-docx non disponible. L'export Word sera désactivé.")

def transcrire_audio(fichier_audio, sortie_dossier, langue="fr", num_speakers=None):
    """Transcrit un fichier audio avec WhisperX et diarisation"""
    try:
        print("──────────────────────────────────────────────")
        print(f"▶️ DÉMARRAGE DE LA TRANSCRIPTION")
        print(f"📂 Fichier audio : {fichier_audio}")
        print(f"🌐 Langue choisie : {langue}")
        print("──────────────────────────────────────────────")

        print("⏳ Chargement du modèle WhisperX (large-v3)...")
        model = whisperx.load_model("large-v3", device="cuda")
        print("✅ Modèle WhisperX chargé.")

        print("📝 Transcription initiale avec WhisperX...")
        result = model.transcribe(fichier_audio, language=langue)
        print("✅ Transcription terminée.")

        if not result.get("segments"):
            print("❌ Aucun segment détecté. Fichier peut-être vide ou incompréhensible.")
            return False

        import gc
        import torch
        from torch import device
        from pyannote.audio import Pipeline

        del model
        torch.cuda.empty_cache()
        gc.collect()

        print("🧹 Mémoire GPU libérée après WhisperX.")

        print("🔊 Début de la diarisation avec pyannote.audio...")
        
        diarize_model = Pipeline.from_pretrained("pyannote/speaker-diarization", use_auth_token=True)
        diarize_model.to(device("cuda"))
        diarization = diarize_model(fichier_audio)
        print("✅ Diarisation terminée.")

        print("📐 Formatage des segments pour alignement...")
        diarize_segments = [
            {"start": turn.start, "end": turn.end, "speaker": speaker}
            for turn, _, speaker in diarization.itertracks(yield_label=True)
        ]
        print(f"🧩 {len(diarize_segments)} segments de locuteurs générés.")

        print("📎 Alignement des locuteurs avec les segments de transcription...")
        # Charger le modèle d'alignement
        model_a, metadata = whisperx.load_align_model(language_code=langue, device="cuda")
        
        # Effectuer l'alignement avec tous les arguments requis
        result_aligned = whisperx.align(result["segments"], model_a, metadata, fichier_audio, "cuda")
        
        # Libérer le modèle d'alignement
        del model_a
        torch.cuda.empty_cache()
        print("✅ Alignement terminé.")

        print("🎯 Attribution des locuteurs aux segments...")
        # Fonction simple pour assigner les locuteurs aux segments de transcription
        def assigner_locuteur_segment(segment_start, segment_end, diarize_segments):
            """Trouve le locuteur principal pour un segment donné"""
            segment_mid = (segment_start + segment_end) / 2
            
            # Chercher le segment de diarisation qui contient le milieu du segment de transcription
            for diag_seg in diarize_segments:
                if diag_seg["start"] <= segment_mid <= diag_seg["end"]:
                    return diag_seg["speaker"]
            
            # Si pas de correspondance exacte, prendre le plus proche
            best_speaker = "Inconnu"
            min_distance = float('inf')
            for diag_seg in diarize_segments:
                distance = min(abs(diag_seg["start"] - segment_mid), abs(diag_seg["end"] - segment_mid))
                if distance < min_distance:
                    min_distance = distance
                    best_speaker = diag_seg["speaker"]
            
            return best_speaker

        # Créer la structure avec locuteurs assignés
        segments_with_speakers = []
        for segment in result_aligned["segments"]:
            segment_start = segment.get("start", 0)
            segment_end = segment.get("end", segment_start)
            
            # Assigner le locuteur à ce segment
            speaker = assigner_locuteur_segment(segment_start, segment_end, diarize_segments)
            
            # Créer le nouveau segment avec le locuteur
            new_segment = segment.copy()
            new_segment["speaker"] = speaker
            segments_with_speakers.append(new_segment)

        # Créer la structure finale compatible avec le reste du code
        result_with_speakers = {"segments": segments_with_speakers}
        
        print("✅ Attribution des locuteurs terminée.")
        
        locuteurs = set(seg.get("speaker", "Inconnu") for seg in result_with_speakers["segments"])
        print(f"🗣️ Locuteurs détectés : {', '.join(sorted(locuteurs))}")

        print("📁 Création du dossier de sortie (si nécessaire)...")
        os.makedirs(sortie_dossier, exist_ok=True)
        base_nom = os.path.splitext(os.path.basename(fichier_audio))[0]

        print("💾 Exportation des fichiers : TXT, Markdown", end="")
        if DOCX_AVAILABLE:
            print(", DOCX")
        else:
            print(" (⚠️ Word non disponible)")

        exporter_txt(result_with_speakers["segments"], sortie_dossier, base_nom)
        exporter_markdown(result_with_speakers["segments"], sortie_dossier, base_nom, fichier_audio)
        
        if DOCX_AVAILABLE:
            exporter_word(result_with_speakers["segments"], sortie_dossier, base_nom, fichier_audio)

        print(f"🎉 Transcription réussie ! Fichiers disponibles dans : {sortie_dossier}")
        print("──────────────────────────────────────────────")
        return True

    except Exception as e:
        print("❌ Une erreur est survenue durant le processus.")
        print(f"🛑 Détail de l'erreur : {e}")
        print("──────────────────────────────────────────────")
        return False

def exporter_txt(segments, sortie_dossier, base_nom):
    """Export en format texte simple"""
    with open(os.path.join(sortie_dossier, base_nom + ".txt"), "w", encoding="utf-8") as f:
        current_speaker = None
        for seg in segments:
            speaker = seg.get('speaker', 'Inconnu')
            text = seg.get('text', '').strip()
            
            if speaker != current_speaker:
                timestamp = format_timestamp(seg.get("start", 0))
                f.write(f"\n[{timestamp}] {speaker} :\n")
                current_speaker = speaker
            
            f.write(f"{text} ")

def exporter_markdown(segments, sortie_dossier, base_nom, fichier_original):
    """Export en format Markdown pour Obsidian"""
    md_path = os.path.join(sortie_dossier, base_nom + ".md")
    
    with open(md_path, "w", encoding="utf-8") as f:
        # En-tête du document
        f.write(f"# Transcription : {base_nom}\n\n")
        f.write(f"**Fichier source :** `{os.path.basename(fichier_original)}`\n")
        f.write(f"**Date de transcription :** {datetime.now().strftime('%d/%m/%Y à %H:%M')}\n")
        f.write(f"**Durée totale :** {format_timestamp(segments[-1].get('end', 0)) if segments else 'Inconnue'}\n\n")
        
        # Résumé des intervenants
        speakers = set(seg.get('speaker', 'Inconnu') for seg in segments)
        f.write(f"**Intervenants :** {', '.join(sorted(speakers))}\n\n")
        f.write("---\n\n")
        
        # Transcription par locuteur
        current_speaker = None
        for seg in segments:
            speaker = seg.get('speaker', 'Inconnu')
            text = seg.get('text', '').strip()
            
            if speaker != current_speaker:
                timestamp = format_timestamp(seg.get('start', 0))
                if current_speaker is not None:
                    f.write("\n\n")
                f.write(f"## {speaker} [{timestamp}]\n\n")
                current_speaker = speaker
            
            f.write(f"{text} ")

def exporter_word(segments, sortie_dossier, base_nom, fichier_original):
    """Export en format Word (.docx)"""
    if not DOCX_AVAILABLE:
        return
        
    try:
        doc = Document()
        
        # Titre principal
        title = doc.add_heading(f'Transcription : {base_nom}', 0)
        
        # Informations du document
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        cells = info_table.rows[0].cells
        cells[0].text = 'Fichier source'
        cells[1].text = os.path.basename(fichier_original)
        
        cells = info_table.rows[1].cells
        cells[0].text = 'Date de transcription'
        cells[1].text = datetime.now().strftime('%d/%m/%Y à %H:%M')
        
        cells = info_table.rows[2].cells
        cells[0].text = 'Durée totale'
        cells[1].text = format_timestamp(segments[-1].get('end', 0)) if segments else 'Inconnue'
        
        speakers = set(seg.get('speaker', 'Inconnu') for seg in segments)
        cells = info_table.rows[3].cells
        cells[0].text = 'Intervenants'
        cells[1].text = ', '.join(sorted(speakers))
        
        doc.add_paragraph()
        
        # Transcription
        doc.add_heading('Transcription', level=1)
        
        current_speaker = None
        current_paragraph = None
        
        for seg in segments:
            speaker = seg.get('speaker', 'Inconnu')
            text = seg.get('text', '').strip()
            
            if speaker != current_speaker:
                timestamp = format_timestamp(seg.get('start', 0))
                if current_speaker is not None:
                    doc.add_paragraph()
                
                # Titre du locuteur avec timestamp
                speaker_heading = doc.add_heading(f'{speaker} [{timestamp}]', level=2)
                current_speaker = speaker
                current_paragraph = doc.add_paragraph()
            
            # Ajouter le texte au paragraphe courant
            if current_paragraph is not None:
                current_paragraph.add_run(text + " ")
        
        # Sauvegarder
        doc_path = os.path.join(sortie_dossier, base_nom + ".docx")
        doc.save(doc_path)
        
    except Exception as e:
        print(f"Erreur lors de l'export Word : {e}")

def format_timestamp(seconds):
    """Convertit les secondes en format MM:SS ou HH:MM:SS"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    else:
        return f"{minutes:02d}:{seconds:02d}"

def choisir_fichier():
    """
    Invite l'utilisateur à entrer le chemin complet du fichier audio via le terminal.
    Vérifie que le chemin existe.
    """
    chemin = input("Entrez le chemin complet vers le fichier audio : ").strip()
    if os.path.isfile(chemin):
        return chemin
    print("❌ Chemin invalide.")
    return ""

def choisir_dossier_sortie():
    """
    Invite l'utilisateur à entrer le chemin complet vers le dossier de sortie via le terminal.
    Vérifie que le dossier existe, sinon utilise ./transcriptions.
    """
    chemin = input("Entrez le chemin complet vers le dossier de sortie : ").strip()
    if os.path.isdir(chemin):
        return chemin
    print("❌ Dossier invalide, utilisation de ./transcriptions")
    return os.path.join(os.getcwd(), "transcriptions")



def afficher_menu():
    """Affiche le menu principal"""
    print("\n" + "="*60)
    print("🎤 TRANSCRIPTION AUDIO AVEC WHISPERX")
    print("="*60)
    print("1. Transcrire un fichier audio")
    print("2. Changer le dossier de sortie")
    print("3. Quitter")
    print("-"*60)

def main():
    print("🎤 Transcription audio avec WhisperX (modèle large-v3)")
    print("📁 Formats de sortie : TXT, Markdown (Obsidian)" + (", Word" if DOCX_AVAILABLE else ""))
    if not DOCX_AVAILABLE:
        print("ℹ️  Pour l'export Word, installez : pip install python-docx")
    
    while True:
        afficher_menu()
        choix = input("\nVotre choix (1-3) : ").strip()

        if choix == "1":
            fichier_audio = choisir_fichier()
            if not fichier_audio:
                print("❌ Aucun fichier sélectionné.")
                continue

            langue = input("Quelle est la langue du fichier audio ? (fr/en) : ").strip().lower()
            if not langue or langue not in ["fr", "en"]:
                print("❌ Langue non reconnue. Utilisation du français par défaut.")
                langue = "fr"
                                  
            nb_speakers_input = input("Combien de personnes parlent dans l'enregistrement ? (laisser vide pour détection auto) : ").strip()
            num_speakers = None
            if nb_speakers_input.isdigit() and int(nb_speakers_input) > 0:
                num_speakers = int(nb_speakers_input)
                print(f"🗣️ Nombre de locuteurs fixé à : {num_speakers}")
            else:
                print("🤖 Détection automatique du nombre de locuteurs")

            sortie_dossier = choisir_dossier_sortie()
            if not sortie_dossier:
                print("❌ Aucun dossier de sortie sélectionné.")
                continue
            print(f"\n🎵 Fichier sélectionné : {os.path.basename(fichier_audio)}")
            
            if transcrire_audio(fichier_audio, sortie_dossier, langue):
                print("\n✅ Transcription réussie !")
                print(f"📁 Fichiers disponibles dans : {sortie_dossier}")

                reponse = input("\nVoulez-vous ouvrir le dossier de sortie ? (o/n) : ").strip().lower()
                if reponse in ['o', 'oui', 'y', 'yes']:
                    try:
                        os.startfile(sortie_dossier)  # Windows
                    except:
                        try:
                            os.system(f'open "{sortie_dossier}"')  # macOS
                        except:
                            os.system(f'xdg-open "{sortie_dossier}"')  # Linux
            else:
                print("❌ Échec de la transcription.")

        elif choix == "2":
            nouveau_dossier = choisir_dossier_sortie()
            if nouveau_dossier:
                sortie_dossier = nouveau_dossier
                print(f"✅ Nouveau dossier de sortie : {sortie_dossier}")
            else:
                print("❌ Aucun dossier sélectionné, le dossier de sortie reste inchangé.")
        
        elif choix == "3":
            print("\n👋 Merci d'avoir utilisé WhisperX Transcriber !")
            break
        
        else:
            print("❌ Choix invalide. Veuillez réessayer.")

if __name__ == "__main__":
    main()
